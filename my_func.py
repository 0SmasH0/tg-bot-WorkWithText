import os
import docx
from docx.shared import Pt
import string
from spellchecker import SpellChecker
import string


def check_text(text_main: str):
    try:
        spell = SpellChecker(language='ru')
        text = text_main.lower().split()
        mistake = list(spell.unknown(text))
        correct = []
        for_mis = mistake[:]

        if not bool(mistake):
            return 'В вашем тексте нет ошибок!!!'
        else:
            for i in for_mis:
                if i[-1] in string.punctuation:
                    k = i[:-1]
                    idx = text.index(i)
                    if k == spell.correction(k):
                        mistake.pop(idx)
                        continue
                    text[idx] = spell.correction(k) + f'{i[-1]}'
                    mistake[idx] = k
                    correct.append(spell.correction(k))
                else:
                    idx = text.index(i.lower())
                    text[idx] = spell.correction(i)
                    correct.append(spell.correction(i))

        if not bool(mistake):
            return 'В вашем тексте нет ошибок!!!'

        res_cor = ''
        idx = 0
        size = len(mistake)
        print(mistake)
        while idx != size:
            res_cor += f"{mistake[idx]} ➡️ {correct[idx]}\n"
            idx += 1

        return ' '.join(text), res_cor
    except (ValueError, TypeError):
        return 'Поддерживается только русский язык'

def decor_text_for_word(text: str, text_size: int, font: str):
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = font
    style.font.size = Pt(text_size)
    doc.add_paragraph(text)
    doc.save('example.docx')
    return


def decor_word_for_word(text, text_size: int, font: str):
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = font
    style.font.size = Pt(text_size)
    doc.add_paragraph(text)
    doc.save('example.docx')
    return
